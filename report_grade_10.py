import pandas as pd
from docx import Document
import os
from datetime import datetime, timedelta
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
import math
import re
import analysis

# Global config
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

class Session:
    def __init__(self, time, teacher):
        self.time = time
        self.teacher = teacher

class SchoolClass:
    def __init__(self, name):
        self.name = name
        self.sessions = {}

    def add_session(self, session_number, session):
        self.sessions[session_number] = session

# Lists and Class Objects
classes = ['10E1','10E2','10E3','10E4']
class_objects = {cls: SchoolClass(cls) for cls in classes}

class_objects['10E1'].add_session('Session 1', Session('1, 2', 'Tim - Nhật'))
class_objects['10E1'].add_session('Session 2', Session('1, 2', 'Vĩ Văn'))
class_objects['10E1'].add_session('Session 3', Session('1, 2', 'Tường Vy'))

class_objects['10E2'].add_session('Session 1', Session('1, 2', 'Kim Ngân'))
class_objects['10E2'].add_session('Session 2', Session('1, 2', 'Tim - Tâm'))
class_objects['10E2'].add_session('Session 3', Session('1, 2', 'Trúc Quỳnh'))

class_objects['10E3'].add_session('Session 1', Session('1, 2', 'Tim - Nhật'))
class_objects['10E3'].add_session('Session 2', Session('1, 2', 'Vĩ Văn'))
class_objects['10E3'].add_session('Session 3', Session('1, 2', 'Kim Quyên'))

class_objects['10E4'].add_session('Session 1', Session('3, 4', 'Kim Ngân')) 
class_objects['10E4'].add_session('Session 2', Session('3, 4', 'Tim - Tâm'))
class_objects['10E4'].add_session('Session 3', Session('3, 4', 'Trúc Quỳnh'))

# Static Data Loading
# Use a more robust way to find files on Vercel
root_path = os.path.dirname(os.path.abspath(__file__))

def load_static_file(name, default_cols=None):
    path = os.path.join(root_path, name)
    try:
        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}. Current Dir Files: {os.listdir(root_path)}")
        return pd.read_excel(path)
    except Exception as e:
        print(f"CRITICAL: Failed to load {name}: {e}")
        return pd.DataFrame(columns=default_cols) if default_cols else pd.DataFrame()

ielts_syllabus = load_static_file('IELTS_syllabus_10.xlsx')
vstep_syllabus = load_static_file('VSTEP_syllabus_10.xlsx')
list_10 = load_static_file('StudentList10.xlsx', default_cols=['User ID'])

# Totals
ielts_lesson_total = 32
vstep_lesson_total = 57


def read_excel(file_path):
    data = pd.read_excel(file_path)
    # Standardize columns by stripping whitespace
    data.columns = [str(c).strip() for c in data.columns]
    return data

def convert_to_exponential(value):
    if pd.isna(value):
        return value
    if isinstance(value, (int, float)) and float(value).is_integer():
        value = int(value)
        exponent = int(math.log10(abs(value)))
        base = value / (10 ** exponent)
        return f"{int(base*10)}E{exponent-1}"
    if isinstance(value, str) and value.isdigit():
        value = int(value)
        exponent = int(math.log10(abs(value)))
        base = value / (10 ** exponent)
        return f"{int(base*10)}E{exponent-1}"
    return value

def clean_data(data, list_df):
    if 'User ID' not in data.columns:
        raise KeyError(f"Cột 'User ID' không tìm thấy trong file dữ liệu tải lên. Các cột tìm thấy: {data.columns.tolist()}")
    if 'User ID' not in list_df.columns:
        raise KeyError(f"Cột 'User ID' không tìm thấy trong file StudentList. Các cột tìm thấy: {list_df.columns.tolist()}")
        
    df = pd.merge(data, list_df, on='User ID')
    df = df[df['Status'] != 'Removed']
    cols_to_select = ['English Class_y','Full Name', 'Study Time', 'Progress','Units(lessons) Passed',
       'Units(lessons) Studied','Type']
    
    if 'Main Class_y' in df.columns:
        cols_to_select.append('Main Class_y')
    elif 'Main Class' in df.columns:
        cols_to_select.append('Main Class')

    df = df[cols_to_select]
    df['Status'] = 'Unknown'
    df['Average_time_per_lesson'] = 'Unknown'
    df.columns = df.columns.str.replace('(', '', regex=False).str.replace(')', '', regex=False).str.replace(' ', '_')
    for column in df.columns:
        df = df.dropna(subset=[column])
    return df

def set_status(class_type, num_user_lesson, ielts_lesson_num, vstep_lesson_num):
    num_week_lesson = 0
    if class_type == "IELTS":
        num_week_lesson = ielts_lesson_num
    elif class_type == "VSTEP":
        num_week_lesson = vstep_lesson_num
    else:
        return "Unknown"

    if num_user_lesson == num_week_lesson:
        return "keep up"
    elif num_user_lesson < num_week_lesson: 
        return "late"
    elif num_user_lesson > num_week_lesson:
        return "far away"
    return "Unknown"

def calculate_average_time_per_lesson_in_minutes(total_time: str, total_lessons: int) -> str:
    if total_lessons != 0 and isinstance(total_time, str):
        try:
            parts = list(map(int, total_time.split(':')))
            if len(parts) >= 2:
                hours, minutes = parts[0], parts[1]
                total_minutes = hours * 60 + minutes
                avg_minutes = round(total_minutes / total_lessons, 1)
                return f"{avg_minutes} phút"
        except Exception:
            return "0 phút"
    return "0 phút"

def add_contribute_to_dataframe(data, ielts_lesson_num, vstep_lesson_num):
    for index, row in data.iterrows():
        status = set_status(row['Type'], row['Unitslessons_Passed'], ielts_lesson_num, vstep_lesson_num)
        average_time = calculate_average_time_per_lesson_in_minutes(row['Study_Time'], row['Unitslessons_Passed'])
        data.at[index, 'Status'] = status
        data.at[index, 'Average_time_per_lesson'] = average_time

# Helper to process feedback
def get_processed_feedback(classes, feedback_path=None):
    if feedback_path is None:
        feedback_path = os.path.join(BASE_DIR, "PCT Teacher Timesheet  (Responses).xlsx")
    
    if not os.path.exists(feedback_path):
        return pd.DataFrame(columns=['Date', 'Class', 'Your_name', 'Comments'])
    
    data_feedback = read_excel(feedback_path)
    data_feedback.columns = data_feedback.columns.str.replace('(', '', regex=False).str.replace(')', '', regex=False).str.replace(' ', '_')
    data_feedback = data_feedback[['Date', 'Class','Your_name', 'Comments']]
    data_feedback = data_feedback.astype(str)
    data_feedback['Date'] = pd.to_datetime(data_feedback['Date'], errors='coerce')
    
    now = datetime.today()
    today = now - timedelta(days=now.isoweekday() % 7)
    one_week_ago = today - timedelta(days=7)
    
    data_feedback = data_feedback[(data_feedback['Date'] >= one_week_ago ) & (data_feedback['Date'] <= today)]
    data_feedback['Class'] = data_feedback['Class'].str.upper()
    data_feedback['Class'] = data_feedback['Class'].apply(convert_to_exponential)
    
    return data_feedback

def create_report(class_name, current_week, vstep_lesson_num, ielts_lesson_num, data, data_feedback, output_dir, 
                  course_vstep="Practical English A2-B2", course_ielts="Practical English A1",
                  total_ielts=32, total_vstep=57):
    template_path = os.path.join(BASE_DIR, 'word_template - Copy.docx')
    if not os.path.exists(template_path):
        template_path = os.path.join(BASE_DIR, 'word_template.docx')
    
    if not os.path.exists(template_path):
        print(f"Error: Template file not found at {template_path}")
        return None
    
    doc = Document(template_path)
    tables = doc.tables

    # Defensive check for class name format
    if len(class_name) < 3:
        class_type = "VSTEP"
    elif class_name[2].upper() == "E":
        class_type = "IELTS"
    else:
        class_type = "VSTEP"

    if class_type == "IELTS":
        num_week_lesson = ielts_lesson_num
        total_lesson = total_ielts
        syllabus = ielts_syllabus[ielts_syllabus['Week'] == current_week].copy()
    else:
        num_week_lesson = vstep_lesson_num
        total_lesson = total_vstep
        syllabus = vstep_syllabus[vstep_syllabus['Week'] == current_week].copy()

    # ... (existing syllabus logic) ...
    syllabus.columns = syllabus.columns.str.replace('(', '', regex=False).str.replace(')', '', regex=False).str.replace(' ', '_')
    if 'Skill_Focus' not in syllabus.columns and 'Skill Focus' in syllabus.columns:
         syllabus.rename(columns={'Skill Focus': 'Skill_Focus'}, inplace=True)
         
    syllabus = syllabus[['Week','Name','Skill_Focus']]
    syllabus = syllabus.dropna()
    syllabus = syllabus.head(3)

    # Add Class Detail
    class_data = []
    if class_name in class_objects:
        school_class = class_objects[class_name]
        for session_number, session in school_class.sessions.items():
            class_data.append([class_name, session_number, session.time, session.teacher])
    
    df_class_sessions = pd.DataFrame(class_data, columns=['Class', 'Session', 'Time', 'Teacher'])
    
    # Fill Session Info
    for row, (_, data_row) in zip(tables[2].rows[1:], df_class_sessions.iterrows()):
        session_number = data_row['Session']
        if session_number in class_objects[class_name].sessions:
            row.cells[2].text = str(class_objects[class_name].sessions[session_number].time)
            row.cells[3].text = str(class_objects[class_name].sessions[session_number].teacher)

    # Fill Syllabus Info
    for row, (_, data_row) in zip(tables[2].rows[1:], syllabus.iterrows()):
        row.cells[4].text = str(data_row['Name'])
        row.cells[5].text = str(data_row['Skill_Focus'])

    df = data[data['English_Class_y'] == class_name].copy()
    
    # Feedback Helper
    def add_feedback_local():
        feedback_class = data_feedback[data_feedback['Class'].str.contains(class_name, na=False)].copy()
        if feedback_class.empty:
            return []
            
        feedback_class = feedback_class.sort_values(by='Date', ascending=True)
        student_info = df[['Full_Name', 'Main_Class_y']].dropna()
        
        counter = 1
        for idx, feedback in feedback_class.iterrows():
            comment_text = str(feedback['Comments'])
            matches = re.findall(r'"([^"]+)"', comment_text)
            for match in matches:
                match_lower = match.strip().lower()
                found_student = student_info[student_info['Full_Name'].str.lower().str.contains(re.escape(match_lower))]
                if not found_student.empty:
                    full_name = found_student.iloc[0]['Full_Name']
                    main_class = found_student.iloc[0]['Main_Class_y']
                    new_text = f'{full_name} {main_class}' 
                    comment_text = comment_text.replace(f'"{match}"', new_text)

            feedback_class.at[idx, 'Comments'] = comment_text
            
            new_row = tables[3].add_row()
            new_row.cells[0].text = str(counter)
            new_row.cells[1].text = str(feedback['Your_name'])
            new_row.cells[2].text = str(comment_text)
            counter += 1
            
        return feedback_class['Comments'].tolist()

    # Fill Header Info
    for paragraph in doc.paragraphs:
        if "Tổng số bài học cho đến thời điểm báo cáo:" in paragraph.text:
            percentage = (num_week_lesson / total_lesson) * 100
            paragraph.text += f"{num_week_lesson}/{total_lesson} ({percentage:.2f}%)"

    first_row = tables[1].rows[0]
    first_row.cells[2].text = class_name

    second_row = tables[1].rows[1]
    if class_type == "IELTS":
        second_row.cells[2].text = course_ielts
    else:
        second_row.cells[2].text = course_vstep

    third_row = tables[1].rows[2]
    third_row.cells[2].text = (datetime.now()).strftime("%d-%m-%Y")

    for row in (first_row.cells[2],second_row.cells[2],third_row.cells[2]):
        for paragraph in row.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    tables[2].rows[1].cells[0].text = str(current_week)
    for paragraph in tables[2].rows[1].cells[0].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Student Lists
    late_students = df[df['Status'] == 'late'].sort_values(by='Unitslessons_Passed', ascending=False)
    keep_up_students = df[df['Status'] == 'keep up'].sort_values(by='Unitslessons_Passed', ascending=False)
    far_away_students = df[df['Status'] == 'far away'].sort_values(by='Unitslessons_Passed', ascending=False)

    def add_students(table, students):
        # Determine which Main Class column name to use (after cleaning spaces to _)
        main_class_col = 'Main_Class_y' if 'Main_Class_y' in students.columns else 'Main_Class'
        columns = ['Full_Name', main_class_col, 'Progress', 'Study_Time', 'Unitslessons_Passed', 'Unitslessons_Studied', 'Average_time_per_lesson']
        counter = 1
        for _, row in students.iterrows():
            try:
                new_row = table.add_row()
                new_row.cells[0].text = str(counter)
                for i, col in enumerate(columns):
                    if i + 1 < len(new_row.cells):
                        val = row[col]
                        new_row.cells[i + 1].text = str(val) if pd.notna(val) else "N/A"
                counter += 1
            except Exception as e:
                print(f"Error adding student row: {e}")
                continue

    add_students(tables[4], far_away_students)    
    add_students(tables[5], keep_up_students)  
    add_students(tables[6], late_students)  

    def add_warning_students(table, students):
        counter = 1
        for _, row in students.iterrows():
            new_row = table.add_row()
            new_row.cells[0].text = str(counter)
            new_row.cells[1].text = str(row['Full_Name'])
            main_class_col = 'Main_Class_y' if 'Main_Class_y' in students.columns else 'Main_Class'
            if main_class_col in row:
                new_row.cells[2].text = str(row[main_class_col])
            else:
                new_row.cells[2].text = "N/A"
            avg_time = str(row['Average_time_per_lesson'])
            new_row.cells[3].text = f"Thời gian trung bình làm bài quá ngắn ({avg_time})"
            counter += 1

    # Filter Warning Students (Avg time < 10 mins)
    warning_students = df[df['Average_time_per_lesson'].apply(
        lambda x: float(x.split()[0]) < 10 if isinstance(x, str) and len(x.split()) > 0 and x.split()[0].replace('.','',1).isdigit() else False
    )].sort_values(by='Average_time_per_lesson', ascending=True)

    if len(tables) > 7:
        add_warning_students(tables[7], warning_students)  

    # Process Feedback Paragraph
    for paragraph in doc.paragraphs:
        if "Tóm tắt tình hình lớp:" in paragraph.text:
            add_feedback_local()

    # Font Styling
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            run.font.size = Pt(12)

    for table in doc.tables[:4]:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                        run.font.size = Pt(11)   

    # Save
    file_path = f'{output_dir}/W{current_week}-PCT-Report-{class_name}.docx'
    try:
        doc.save(file_path)
        return file_path
    except Exception as e:
        print(f"Error saving {file_path}: {e}")
        return None

def generate_grade_10_reports(current_week, vstep_lesson_num, ielts_lesson_num, data_file_path, output_dir=None, 
                              target_classes_list=None, course_vstep="Practical English A2-B2", course_ielts="Practical English A1",
                              total_ielts=32, total_vstep=57, timesheet_path=None):
    if output_dir is None:
        output_dir = os.path.join(BASE_DIR, 'Grade_10', f'Grade_10_Week {current_week}')
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Use default classes if none provided
    if not target_classes_list:
        target_classes_list = classes  # Fallback to global if needed, or pass from app

    # Process List_10
    if not list_10.empty:
        list_10['English Class'] = list_10['English Class'].apply(convert_to_exponential)

    # Load Data
    excel_file = pd.read_excel(data_file_path)
    
    # Clean and Process
    data = clean_data(excel_file, list_10)
    add_contribute_to_dataframe(data, ielts_lesson_num, vstep_lesson_num)
    
    # Get Feedback Data
    data_feedback = get_processed_feedback(target_classes_list, feedback_path=timesheet_path)

    generated_files = []
    
    # Calculate Stats
    stats = analysis.calculate_stats(data)

    for class_name in target_classes_list:
        # Check if class name format is valid (needed for class_objects lookup or simple object creation)
        if class_name not in class_objects:
            # Dynamically create if config has new class names? 
            # For now, just warn or skip, or create simple object
            print(f"Warning: Class {class_name} not found in session config. Report will lack session details.")
            # We can still generate report, just without session info if it depends on class_objects
        
        path = create_report(class_name, current_week, vstep_lesson_num, ielts_lesson_num, data, data_feedback, output_dir, course_vstep, course_ielts, total_ielts, total_vstep)
        if path:
            generated_files.append(path)
            
    return generated_files, stats
